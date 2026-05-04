"""
Generate Vigilent Optimization Model - Client Walkthrough Document
===================================================================
Produces a .docx formatted to match Ayush's BEACNomics Curve Research Document style.

Usage:
    python3 generate_client_walkthrough_doc.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_PATH = "output/Vigilent_Optimization_Model_Walkthrough.docx"

# --- Vigilent brand colors ---
NAVY = RGBColor(0x1B, 0x28, 0x5B)
BLUE = RGBColor(0x10, 0x75, 0xE8)
BLACK = RGBColor(0x00, 0x00, 0x00)
HEADER_FILL = "d7e7f8"  # light blue table header
ALT_ROW_FILL = "f2f7fc"  # alternating row shading


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Source Sans Pro"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    run.font.name = "Source Sans Pro"
    p.space_before = Pt(12)


def add_subsection_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.font.name = "Source Sans Pro"
    p.space_before = Pt(8)


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Source Sans Pro"
    return p


def add_body_blue(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE
    run.font.name = "Source Sans Pro"
    return p


def add_summary_table(doc, rows):
    """2-column key-value table (like Shared Baseline Assumptions)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        r0 = c0.paragraphs[0].add_run(key)
        r0.font.size = Pt(12)
        r0.font.name = "Source Sans Pro"
        r0.bold = True
        r1 = c1.paragraphs[0].add_run(val)
        r1.font.size = Pt(12)
        r1.font.name = "Source Sans Pro"
    doc.add_paragraph()  # spacing


def add_result_table(doc, rows):
    """2-column result summary (like Abatement Cost / Abatement Potential)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        r0 = c0.paragraphs[0].add_run(key)
        r0.font.size = Pt(12)
        r0.font.name = "Source Sans Pro"
        r0.bold = True
        r1 = c1.paragraphs[0].add_run(val)
        r1.font.size = Pt(12)
        r1.font.name = "Source Sans Pro"
        r1.bold = True
    doc.add_paragraph()


def add_step_table(doc, header_text, steps):
    """3-column step-by-step calculation table matching reference format."""
    table = doc.add_table(rows=len(steps) + 2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Merged header row
    hdr = table.cell(0, 0)
    hdr.merge(table.cell(0, 1)).merge(table.cell(0, 2))
    r = hdr.paragraphs[0].add_run(header_text)
    r.font.size = Pt(12)
    r.font.name = "Source Sans Pro"
    r.bold = True
    set_cell_shading(hdr, HEADER_FILL)

    # Column headers
    for j, label in enumerate(["Step", "Formula/Inputs", "Result"]):
        cell = table.cell(1, j)
        r = cell.paragraphs[0].add_run(label)
        r.font.size = Pt(12)
        r.font.name = "Source Sans Pro"
        r.bold = True

    # Data rows
    for i, (step, formula, result) in enumerate(steps):
        for j, text in enumerate([step, formula, result]):
            cell = table.cell(i + 2, j)
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(12)
            r.font.name = "Source Sans Pro"

    doc.add_paragraph()


def add_scoring_table(doc, header, rows):
    """Multi-column table with header shading."""
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows) + 1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for j, label in enumerate(header):
        cell = table.cell(0, j)
        r = cell.paragraphs[0].add_run(label)
        r.font.size = Pt(12)
        r.font.name = "Source Sans Pro"
        r.bold = True
        set_cell_shading(cell, HEADER_FILL)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, text in enumerate(row_data):
            cell = table.cell(i + 1, j)
            r = cell.paragraphs[0].add_run(str(text))
            r.font.size = Pt(12)
            r.font.name = "Source Sans Pro"
            if i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_FILL)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Source Sans Pro"
    font.size = Pt(12)

    # ── TITLE ──
    add_title(doc, "Vigilent - Optimization Model Walkthrough")
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 1. DOCUMENT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "Document Overview")
    doc.add_paragraph()
    add_body(doc,
        "This document provides a complete, step-by-step explanation of the Vigilent "
        "Data Center Optimization Model. It covers how input parameters are transformed "
        "into a composite score, how financial metrics (ROI, payback period, savings per MW) "
        "are calculated, and how environmental justice impact is assessed. All formulas, "
        "assumptions, data sources, and scoring weights are documented so the client can "
        "independently verify, adjust, and extend the model."
    )
    doc.add_paragraph()
    add_body(doc,
        "Three outputs drive the model:"
    )
    add_body(doc,
        "Composite Score (0\u2013100) \u2014 a weighted index that ranks how well a data center "
        "fits the Vigilent value proposition. Higher scores indicate stronger ROI and operational impact."
    )
    add_body(doc,
        "Financial Metrics \u2014 estimated annual savings, savings per MW, payback period, "
        "and OPEX impact percentage."
    )
    add_body(doc,
        "Environmental Justice Impact \u2014 CO\u2082 avoided, water saved, grid relief, and "
        "community marginalization indicators tied to the data center's location."
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 2. SHARED BASELINE ASSUMPTIONS
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "Shared Baseline Assumptions")
    doc.add_paragraph()
    add_body(doc,
        "The following inputs are common across all data centers when site-specific data is unavailable. "
        "These represent industry-standard estimates and can be overridden with real facility data."
    )
    doc.add_paragraph()

    add_summary_table(doc, [
        ("Baseline PUE", "1.55 (Uptime Institute 2023 global average)"),
        ("Load Growth Rate", "10% annual (industry consensus estimate)"),
        ("Energy % of OPEX", "40% (Gartner, McKinsey typical for large DCs)"),
        ("Analysis Horizon", "1 year (snapshot model)"),
        ("Vigilent Investment Cost", "$1,500,000 (standard offering)"),
        ("Energy Reduction with Vigilent", "10% (standard offering)"),
        ("Water Reduction with Vigilent", "5% (standard offering)"),
    ])

    add_subsection_header(doc, "Data Sources")
    add_body(doc, "Uptime Institute \u2014 Global PUE average (2023 annual survey)")
    add_body(doc, "EIA (U.S. Energy Information Administration) \u2014 State-level commercial electricity rates")
    add_body(doc, "Gartner / McKinsey \u2014 Energy as % of data center operating costs")
    add_body(doc, "EPA eGRID2022 \u2014 Grid emission factors and fuel mix by subregion")
    add_body(doc, "U.S. Census Bureau ACS 2023 \u2014 Poverty rate, median income, demographics")
    add_body(doc, "DOE LEAD Tool / EIA RECS 2020 \u2014 Energy burden by state")
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 3. MODEL COMPONENTS \u2014 DETAILED METHODOLOGY
    # ══════════════════════════════════════════════════════════════════════
    add_title(doc, "Model Components \u2014 Detailed Methodology")
    doc.add_paragraph()

    # ── 3.1 ENERGY & FINANCIAL CALCULATIONS ──
    add_section_header(doc, "Energy & Financial Calculations")
    doc.add_paragraph()

    add_result_table(doc, [
        ("Outputs Produced", "Annual Savings, Savings/MW, Payback Period, OPEX Impact %"),
        ("Input Parameters", "DC Size (MW), Baseline PUE, Electricity Price, Load Growth, Energy % OPEX"),
    ])

    add_step_table(doc, "How Annual Energy Cost is Calculated", [
        ("Calculate total power draw",
         "DC Size (MW) \u00d7 Baseline PUE",
         "Total Power (MW)"),
        ("Convert to annual energy",
         "Total Power \u00d7 1,000 \u00d7 8,760 hrs/yr \u00d7 (1 + Load Growth Rate)",
         "Annual Energy (kWh)"),
        ("Calculate annual energy cost",
         "Annual Energy (kWh) \u00d7 Electricity Price ($/kWh)",
         "Annual Energy Cost ($)"),
    ])

    add_step_table(doc, "How Financial Metrics are Calculated", [
        ("Estimated annual savings",
         "Annual Energy Cost \u00d7 Energy Reduction % (10%)",
         "Estimated Savings ($)"),
        ("Savings per MW",
         "Estimated Savings / DC Size (MW)",
         "$/MW"),
        ("Payback period",
         "Investment Cost ($1,500,000) / Estimated Savings",
         "Years"),
        ("OPEX impact",
         "Energy Reduction % \u00d7 Energy % of OPEX",
         "% of operating costs reduced"),
    ])

    add_subsection_header(doc, "Worked Example \u2014 20 MW Data Center in Virginia")
    doc.add_paragraph()

    add_step_table(doc, "Example: 20 MW DC, PUE 1.55, Electricity $0.0973/kWh, 10% Load Growth", [
        ("Total power draw",
         "20 MW \u00d7 1.55",
         "31 MW"),
        ("Annual energy consumption",
         "31 MW \u00d7 1,000 \u00d7 8,760 \u00d7 1.10",
         "298,452,000 kWh"),
        ("Annual energy cost",
         "298,452,000 kWh \u00d7 $0.0973/kWh",
         "$29,039,380"),
        ("Estimated annual savings (10%)",
         "$29,039,380 \u00d7 0.10",
         "$2,903,938"),
        ("Savings per MW",
         "$2,903,938 / 20 MW",
         "$145,197/MW"),
        ("Payback period",
         "$1,500,000 / $2,903,938",
         "0.52 years (6.2 months)"),
        ("OPEX impact",
         "0.10 \u00d7 0.40",
         "4.0% of OPEX reduced"),
    ])

    add_subsection_header(doc, "Key Assumptions")
    add_body(doc,
        "PUE is applied as a constant multiplier. In practice, PUE varies by season and load. "
        "Site-specific PUE data will improve accuracy."
    )
    add_body(doc,
        "Electricity price uses state-level EIA commercial rates. Actual negotiated rates "
        "may differ significantly, especially for large consumers."
    )
    add_body(doc,
        "The model uses a single-year snapshot. No multi-year NPV, depreciation, or discount rate "
        "is applied."
    )
    doc.add_paragraph()

    # ── 3.2 COMPOSITE SCORING MODEL ──
    add_section_header(doc, "Composite Scoring Model")
    doc.add_paragraph()

    add_body(doc,
        "The composite score ranks data centers from 0 to 100 based on five weighted factors. "
        "Each factor is normalized to a 0\u2013100 scale before weighting. The final score indicates "
        "how strongly a data center benefits from Vigilent's solution."
    )
    doc.add_paragraph()

    add_result_table(doc, [
        ("Score Range", "0\u2013100"),
        ("Classification", "Excellent (75\u2013100), Good (50\u201374), Moderate (25\u201349), Low (0\u201324)"),
    ])

    add_scoring_table(doc,
        ["Factor", "Weight", "Min", "Max", "Direction", "Why It Matters"],
        [
            ("Savings per MW", "35%", "$0", "$300,000", "Higher = better",
             "Core ROI metric. Larger savings/MW = stronger value proposition"),
            ("Payback Period", "25%", "0 yr", "5 yr", "INVERTED (shorter = better)",
             "Speed of cost recovery. Sub-1-year payback is ideal"),
            ("OPEX Impact %", "20%", "0%", "10%", "Higher = better",
             "Share of operating costs reduced. Drives CFO-level interest"),
            ("Water Savings %", "10%", "0%", "8%", "Higher = better",
             "Sustainability metric. Relevant for ESG reporting"),
            ("Load Growth Rate", "10%", "0%", "15%", "Higher = better",
             "Growing facilities benefit more from efficiency gains"),
        ]
    )

    add_step_table(doc, "How Each Factor is Normalized (0\u2013100 Scale)", [
        ("Standard normalization (4 factors)",
         "Score = ((Value \u2013 Min) / (Max \u2013 Min)) \u00d7 100, clamped to [0, 100]",
         "0\u2013100"),
        ("Inverted normalization (Payback only)",
         "Score = (1 \u2013 (Value \u2013 Min) / (Max \u2013 Min)) \u00d7 100, clamped to [0, 100]",
         "0\u2013100"),
        ("Composite score",
         "0.35 \u00d7 Savings + 0.25 \u00d7 Payback + 0.20 \u00d7 OPEX + 0.10 \u00d7 Water + 0.10 \u00d7 Growth",
         "0\u2013100"),
    ])

    add_subsection_header(doc, "Worked Example \u2014 Scoring the 20 MW Virginia DC")
    doc.add_paragraph()

    add_step_table(doc, "Example: Savings/MW $145,197, Payback 0.52 yr, OPEX 4%, Water 5%, Growth 10%", [
        ("Savings per MW score",
         "($145,197 \u2013 $0) / ($300,000 \u2013 $0) \u00d7 100",
         "48.4"),
        ("Payback period score (inverted)",
         "(1 \u2013 (0.52 \u2013 0) / (5.0 \u2013 0)) \u00d7 100",
         "89.6"),
        ("OPEX impact score",
         "(0.04 \u2013 0) / (0.10 \u2013 0) \u00d7 100",
         "40.0"),
        ("Water savings score",
         "(0.05 \u2013 0) / (0.08 \u2013 0) \u00d7 100",
         "62.5"),
        ("Load growth score",
         "(0.10 \u2013 0) / (0.15 \u2013 0) \u00d7 100",
         "66.7"),
        ("Composite score",
         "48.4\u00d70.35 + 89.6\u00d70.25 + 40.0\u00d70.20 + 62.5\u00d70.10 + 66.7\u00d70.10",
         "60.3 (Good)"),
    ])

    add_subsection_header(doc, "Key Assumptions")
    add_body(doc,
        "Scoring weights (35/25/20/10/10) reflect Vigilent's sales prioritization: financial ROI "
        "(savings per MW and payback) account for 60% of the score."
    )
    add_body(doc,
        "Factor ranges (Min/Max) define the normalization window. Values beyond Max still score 100; "
        "values below Min score 0. Ranges are based on observed industry data."
    )
    add_body(doc,
        "Payback period uses inverted scoring because shorter payback = better outcome. A 0-year "
        "payback scores 100; a 5-year payback scores 0."
    )
    doc.add_paragraph()

    # ── 3.3 ENVIRONMENTAL JUSTICE IMPACT ──
    add_section_header(doc, "Environmental Justice (EJ) Impact Model")
    doc.add_paragraph()

    add_body(doc,
        "The EJ model quantifies the environmental and community impact of Vigilent's energy "
        "reduction at each data center location. It uses EPA eGRID data to convert energy savings "
        "into carbon avoided, water saved, and grid relief, then contextualizes these against "
        "local marginalization indicators."
    )
    doc.add_paragraph()

    add_result_table(doc, [
        ("Outputs", "CO\u2082 Avoided (MT/yr), Water Saved (gal), Grid Relief %, Demographic Index"),
        ("Input", "DC Size, PUE, Load Growth, Energy Reduction %, Zip Code"),
    ])

    add_step_table(doc, "How Carbon Avoided is Calculated", [
        ("Total annual energy",
         "DC Size (MW) \u00d7 Baseline PUE \u00d7 8,760 hrs \u00d7 (1 + Load Growth)",
         "Total Energy (MWh)"),
        ("Energy saved by Vigilent",
         "Total Energy (MWh) \u00d7 Energy Reduction %",
         "Energy Saved (MWh)"),
        ("Resolve grid emission factor",
         "Zip code \u2192 State \u2192 eGRID subregion \u2192 CO\u2082 rate (lbs/MWh)",
         "lbs CO\u2082/MWh"),
        ("Carbon avoided (lbs)",
         "Energy Saved (MWh) \u00d7 CO\u2082 Rate (lbs/MWh)",
         "lbs CO\u2082/yr"),
        ("Carbon avoided (metric tons)",
         "lbs CO\u2082 / 2,204.62",
         "MT CO\u2082/yr"),
    ])

    add_step_table(doc, "How Water Saved is Calculated", [
        ("Determine fuel mix",
         "eGRID subregion \u2192 % coal, gas, nuclear, hydro, wind, solar",
         "Fuel mix (%)"),
        ("Calculate water intensity",
         "\u03a3(Fuel % \u00d7 Water per Fuel Type)",
         "gal/MWh"),
        ("Water saved",
         "Energy Saved (MWh) \u00d7 Water Intensity (gal/MWh)",
         "gallons/yr"),
    ])

    add_subsection_header(doc, "Water Intensity by Fuel Type")
    add_scoring_table(doc,
        ["Fuel Type", "Water Intensity (gal/MWh)", "Source"],
        [
            ("Coal", "12,000", "NREL / USGS / Macknick et al. 2012"),
            ("Natural Gas", "2,800", "NREL / USGS / Macknick et al. 2012"),
            ("Nuclear", "13,000", "NREL / USGS / Macknick et al. 2012"),
            ("Hydro", "0 (no consumptive cooling)", "NREL"),
            ("Wind", "0", "NREL"),
            ("Solar (PV)", "26 (panel cleaning)", "NREL"),
            ("Other/Biomass", "5,000", "NREL"),
        ]
    )

    add_step_table(doc, "How Grid Relief is Calculated", [
        ("MW freed by Vigilent",
         "DC Size (MW) \u00d7 Baseline PUE \u00d7 Energy Reduction %",
         "MW freed"),
        ("Grid relief percentage",
         "(MW Freed / State Peak Demand MW) \u00d7 100",
         "% of state grid"),
    ])

    add_step_table(doc, "EPA Equivalency Conversions", [
        ("Cars equivalent",
         "CO\u2082 Avoided (lbs) / 10,141 lbs per car per year",
         "cars/yr"),
        ("Homes equivalent",
         "Energy Saved (MWh) \u00d7 1,000 / 10,500 kWh per home per year",
         "homes/yr"),
        ("Trees equivalent",
         "CO\u2082 Avoided (lbs) / 48 lbs per tree per year",
         "trees/yr"),
        ("Olympic pools equivalent",
         "Water Saved (gal) / 660,000 gal per pool",
         "pools"),
    ])

    add_subsection_header(doc, "Demographic Index (Marginalization Indicator)")
    doc.add_paragraph()
    add_body(doc,
        "The Demographic Index approximates EPA's EJScreen methodology to assess community "
        "vulnerability near each data center."
    )
    doc.add_paragraph()

    add_step_table(doc, "How Demographic Index is Calculated", [
        ("Low income approximation",
         "min(State Poverty Rate \u00d7 2, 60%)",
         "% low income"),
        ("Demographic Index",
         "(Low Income Approx + People of Color %) / 2",
         "0\u2013100 index"),
        ("National average reference",
         "(min(12.4% \u00d7 2, 60%) + 39.6%) / 2",
         "32.2 (national avg)"),
    ])

    add_subsection_header(doc, "Key Assumptions")
    add_body(doc,
        "All EJ metrics use state-level aggregates. Zip-code-level data would provide more "
        "granular community impact estimates."
    )
    add_body(doc,
        "Grid emission factors are from EPA eGRID2022 and assume a static fuel mix. "
        "Ongoing renewable energy transition is not modeled."
    )
    add_body(doc,
        "Water intensity depends on power plant cooling type (air vs. wet). The model uses "
        "fuel-type-based national averages."
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 4. INPUT PARAMETERS REFERENCE
    # ══════════════════════════════════════════════════════════════════════
    add_title(doc, "Input Parameters Reference")
    doc.add_paragraph()

    add_body(doc,
        "The table below lists every input parameter, its valid range, default value, "
        "and data source. Parameters marked REAL come from actual facility or government data. "
        "Parameters marked ESTIMATED use industry averages and should be replaced with "
        "site-specific data when available."
    )
    doc.add_paragraph()

    add_scoring_table(doc,
        ["Parameter", "Range", "Default", "Source", "Data Quality"],
        [
            ("DC Size (MW)", "1\u2013200 MW", "20 MW", "Facility records / CSV", "REAL"),
            ("Baseline PUE", "1.0\u20132.5", "1.55", "Uptime Institute 2023", "ESTIMATED"),
            ("Electricity Price ($/kWh)", "$0.01\u2013$0.50", "$0.10", "EIA state commercial rates", "REAL (state avg)"),
            ("Load Growth Rate", "0\u201330%", "10%", "Industry consensus", "ESTIMATED"),
            ("Energy % of OPEX", "5\u201380%", "40%", "Gartner / McKinsey", "ESTIMATED"),
            ("Investment Cost ($)", "$10K\u2013$50M", "$1,500,000", "Vigilent standard offering", "PARAMETER"),
            ("Energy Reduction %", "1\u201340%", "10%", "Vigilent standard offering", "PARAMETER"),
            ("Water Reduction %", "1\u201315%", "5%", "Vigilent standard offering", "PARAMETER"),
            ("Analysis Years", "1\u201310", "1", "User selection", "PARAMETER"),
        ]
    )

    # ══════════════════════════════════════════════════════════════════════
    # 5. HOW TO ADJUST THE MODEL
    # ══════════════════════════════════════════════════════════════════════
    add_title(doc, "How to Adjust Inputs and Understand Impact")
    doc.add_paragraph()

    add_body(doc,
        "The model is designed so that any input parameter can be adjusted to reflect "
        "site-specific data. Below is a guide to which inputs have the largest impact on results."
    )
    doc.add_paragraph()

    add_scoring_table(doc,
        ["Input Changed", "Primary Effect", "Impact on Score", "Sensitivity"],
        [
            ("Electricity Price \u2191", "Savings increase", "Score increases significantly",
             "HIGH \u2014 directly scales all financial metrics"),
            ("DC Size \u2191", "Total savings increase, savings/MW stays similar",
             "Score stays similar (ratio-based)", "LOW on score, HIGH on absolute $"),
            ("Baseline PUE \u2191", "More energy consumed, more savings possible",
             "Score increases moderately", "MEDIUM"),
            ("Load Growth Rate \u2191", "More energy consumed + direct factor score increase",
             "Score increases", "MEDIUM"),
            ("Energy % OPEX \u2191", "OPEX impact factor increases",
             "Score increases moderately", "MEDIUM"),
            ("Investment Cost \u2191", "Payback period lengthens",
             "Score decreases (payback factor worsens)", "HIGH"),
            ("Energy Reduction % \u2191", "All savings increase, payback shortens",
             "Score increases significantly", "HIGH"),
        ]
    )

    add_subsection_header(doc, "Sensitivity Analysis")
    add_body(doc,
        "The model includes a built-in sensitivity analysis that tests each estimated input "
        "at low and high bounds:"
    )
    doc.add_paragraph()

    add_scoring_table(doc,
        ["Estimated Input", "Low Bound", "Default", "High Bound"],
        [
            ("Baseline PUE", "1.20", "1.55", "1.80"),
            ("Load Growth Rate", "5%", "10%", "15%"),
            ("Energy % of OPEX", "30%", "40%", "50%"),
        ]
    )

    add_body(doc,
        "For each data center, the model computes the best-case score (all favorable bounds), "
        "worst-case score (all unfavorable bounds), and the resulting score range. "
        "A score range greater than 15 points is flagged as HIGH priority for real data collection."
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 6. LIMITATIONS & CONSTRAINTS
    # ══════════════════════════════════════════════════════════════════════
    add_title(doc, "Limitations, Constraints & Considerations")
    doc.add_paragraph()

    add_subsection_header(doc, "Model Limitations")
    add_body(doc,
        "Single-year snapshot: The model does not account for multi-year depreciation, "
        "discount rates, or escalating electricity prices. Payback period is a simple ratio, "
        "not an NPV-based calculation."
    )
    add_body(doc,
        "State-level granularity: Electricity prices, EJ metrics, and grid emission factors "
        "are state-level averages. Actual site-level data may differ substantially."
    )
    add_body(doc,
        "Static grid assumption: eGRID emission factors reflect a point-in-time fuel mix. "
        "As states add renewable capacity, actual carbon impact will change."
    )
    add_body(doc,
        "Uniform PUE assumption: All data centers without site-specific PUE use the same "
        "industry average (1.55). Real PUE ranges from 1.1 (hyperscale) to 2.5+ (legacy)."
    )
    doc.add_paragraph()

    add_subsection_header(doc, "Operational Constraints")
    add_body(doc,
        "Minimum facility size: The model is most relevant for data centers \u2265 1 MW. Smaller "
        "facilities may not justify the standard investment cost."
    )
    add_body(doc,
        "Cooling infrastructure compatibility: Vigilent's AI cooling optimization requires "
        "existing mechanical cooling systems. Facilities using only free-air cooling may "
        "see reduced benefit."
    )
    add_body(doc,
        "Regulatory constraints: States with active data center energy regulations "
        "(CA, IL, VA, TX) may impose additional compliance requirements that affect "
        "implementation timelines."
    )
    doc.add_paragraph()

    add_subsection_header(doc, "Financial Constraints")
    add_body(doc,
        "The model assumes a fixed investment cost of $1,500,000. Actual implementation "
        "costs vary based on facility size, complexity, and existing infrastructure. "
        "The investment cost parameter can be adjusted in the model to reflect site-specific quotes."
    )
    add_body(doc,
        "Payback periods exceeding 5 years score 0 on the payback factor. This threshold "
        "can be adjusted by modifying the scoring configuration."
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 7. FORMULA REFERENCE CARD
    # ══════════════════════════════════════════════════════════════════════
    add_title(doc, "Formula Reference Card")
    doc.add_paragraph()

    add_scoring_table(doc,
        ["Calculation", "Formula"],
        [
            ("Total Power (MW)", "DC Size \u00d7 Baseline PUE"),
            ("Annual Energy (kWh)", "Total Power \u00d7 1,000 \u00d7 8,760 \u00d7 (1 + Load Growth)"),
            ("Annual Energy Cost ($)", "Annual Energy \u00d7 Electricity Price"),
            ("Estimated Annual Savings ($)", "Annual Energy Cost \u00d7 Energy Reduction %"),
            ("Savings per MW ($/MW)", "Estimated Savings / DC Size"),
            ("Payback Period (yr)", "Investment Cost / Estimated Savings"),
            ("OPEX Impact (%)", "Energy Reduction % \u00d7 Energy % of OPEX"),
            ("CO\u2082 Avoided (MT)", "Energy Saved (MWh) \u00d7 CO\u2082 Rate / 2,204.62"),
            ("Water Saved (gal)", "Energy Saved (MWh) \u00d7 Water Intensity"),
            ("Grid Relief (%)", "(MW Freed / State Peak Demand) \u00d7 100"),
            ("Factor Score (standard)", "((Value \u2013 Min) / (Max \u2013 Min)) \u00d7 100"),
            ("Factor Score (inverted)", "(1 \u2013 (Value \u2013 Min) / (Max \u2013 Min)) \u00d7 100"),
            ("Composite Score",
             "0.35\u00d7Savings + 0.25\u00d7Payback + 0.20\u00d7OPEX + 0.10\u00d7Water + 0.10\u00d7Growth"),
        ]
    )

    # ── SAVE ──
    import os
    os.makedirs("output", exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
